"""Tests for DOCX table parsing via DocxWithTablesLoader."""

import tempfile
import os
from docx import Document as DocxDocument
from services.docx_loader import DocxWithTablesLoader


def _create_test_docx(path: str):
    """Create a .docx file with paragraphs and a table."""
    doc = DocxDocument()
    doc.add_paragraph("Introduction paragraph before the table.")

    table = doc.add_table(rows=3, cols=3)
    data = [
        ["Name", "Age", "City"],
        ["Alice", "30", "New York"],
        ["Bob", "25", "London"],
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    doc.add_paragraph("Conclusion paragraph after the table.")
    doc.save(path)


def test_docx_loader_extracts_tables():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        _create_test_docx(path)
        docs = DocxWithTablesLoader(path).load()

        assert len(docs) == 1
        text = docs[0].page_content

        # Paragraphs are present
        assert "Introduction paragraph" in text
        assert "Conclusion paragraph" in text

        # Table data is present as Markdown
        assert "| Name | Age | City |" in text
        assert "| --- | --- | --- |" in text
        assert "| Alice | 30 | New York |" in text
        assert "| Bob | 25 | London |" in text

        # Order is preserved: intro before table, table before conclusion
        intro_pos = text.index("Introduction")
        table_pos = text.index("| Name |")
        conclusion_pos = text.index("Conclusion")
        assert intro_pos < table_pos < conclusion_pos
    finally:
        os.unlink(path)


def test_docx_loader_no_tables():
    """A docx with no tables should still load paragraphs fine."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = DocxDocument()
        doc.add_paragraph("Just a simple paragraph.")
        doc.add_paragraph("Another paragraph.")
        doc.save(path)

        docs = DocxWithTablesLoader(path).load()
        text = docs[0].page_content

        assert "Just a simple paragraph." in text
        assert "Another paragraph." in text
    finally:
        os.unlink(path)


def test_docx_loader_metadata_source():
    """Metadata should contain the filename."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = DocxDocument()
        doc.add_paragraph("Hello")
        doc.save(path)

        docs = DocxWithTablesLoader(path).load()
        assert docs[0].metadata["source"] == os.path.basename(path)
    finally:
        os.unlink(path)
