"""
Custom DOCX loader that preserves table content as Markdown tables.
Replaces Docx2txtLoader which silently drops all table data.
"""

from pathlib import Path
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from langchain_core.documents import Document


class DocxWithTablesLoader:
    """Loads .docx files, extracting both paragraphs and tables in document order."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def _table_to_markdown(self, table) -> str:
        """Convert a docx Table object to a Markdown-formatted table string."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) >= 1:
            # Insert header separator after the first row
            col_count = len(table.rows[0].cells)
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def load(self) -> list[Document]:
        """Load the docx file and return a list of LangChain Document objects."""
        doc = DocxDocument(self.file_path)
        parts = []

        for element in doc.element.body:
            if element.tag == qn("w:p"):
                # Paragraph
                text = element.text
                if text and text.strip():
                    parts.append(text.strip())
            elif element.tag == qn("w:tbl"):
                # Table — find the matching Table object
                for table in doc.tables:
                    if table._tbl is element:
                        md_table = self._table_to_markdown(table)
                        if md_table:
                            parts.append(md_table)
                        break

        full_text = "\n\n".join(parts)
        metadata = {"source": Path(self.file_path).name}
        return [Document(page_content=full_text, metadata=metadata)]
